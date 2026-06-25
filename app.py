import nltk
from nltk.corpus import stopwords
from nltk.tokenize import PunktSentenceTokenizer, word_tokenize

nltk.data.path.append(r"C:\Users\TANUSHRI\AppData\Roaming\nltk_data")

from nltk.tokenize import word_tokenize
from flask import Flask, render_template, request, jsonify, session
import os
import fitz
import pandas as pd
import spacy
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import pdfplumber
import docx
import re
import nltk
from nltk.tokenize import PunktSentenceTokenizer
from nltk.corpus import stopwords
from sentence_transformers import SentenceTransformer
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from langdetect import detect
from werkzeug.utils import secure_filename  # For secure_filename function
import uuid
import time
import requests
import json
from module3_backend import run_analysis
import plotly
import time
from difflib import SequenceMatcher
import json
import subprocess
import threading
from job_fetcher import fetch_and_store_jobs
from flask import Flask, render_template, request, jsonify, session, redirect, url_for, flash
from werkzeug.security import generate_password_hash, check_password_hash
from auth_database import create_users_table, get_user_by_email, create_user

from dotenv import load_dotenv

load_dotenv()

import google.generativeai as genai

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, "data")
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")

# genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# gemini_model = genai.GenerativeModel(
#     model_name="models/gemini-2.5-flash"
# )

app = Flask(__name__)
app.secret_key = "super-secret-key-change-this"
create_users_table()
app.config['UPLOAD_FOLDER'] = 'uploads'

os.makedirs('uploads', exist_ok=True)

def run_job_fetcher(role):
    try:
        fetch_and_store_jobs(role)
    except Exception as e:
        print("Job fetcher failed (ignored):", e)

# -------- LOAD MODELS & DATA (RUN ONCE) --------
# nlp = spacy.load("en_core_web_sm")

try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    print("⚠️ spaCy model 'en_core_web_sm' not found. Install with: python -m spacy download en_core_web_sm")
    print("Module 3 will still work (uses sentence-transformers instead)")
    nlp = None  # We'll use it only where needed

# GLOBAL MODEL (load once)
embedder = SentenceTransformer("all-MiniLM-L6-v2")

# job_role_df = pd.read_excel("data/job role and skills.xlsx")
app.config["UPLOAD_FOLDER"] = UPLOAD_DIR
os.makedirs(UPLOAD_DIR, exist_ok=True)

job_role_df = pd.read_excel(os.path.join(DATA_DIR, "job role and skills.xlsx"))

# def load_skill_keywords():
#     doc = Document("data/Skill dataset.docx")
#     skills = []
#     for p in doc.paragraphs:
#         skills.extend(p.text.lower().split())
#     return list(set(skills))

# def load_skill_keywords():
#     doc = Document(os.path.join(DATA_DIR, "Skills dataset.docx"))
#     skills = []
#     for p in doc.paragraphs:
#         # skills.extend(p.text.lower().split())
#         skills.append(p.text.strip().lower())
#     return list(set(skills))

def load_skill_keywords():
    doc = Document(os.path.join(DATA_DIR, "Skills dataset.docx"))
    skills = []

    for p in doc.paragraphs:
        line = p.text.strip().lower()

        if line:
            parts = re.split(r"[,\n]|\s+(?=[a-zA-Z])", line)

            for skill in parts:
                skill = skill.strip()

                # 🔥 FILTER NOISE
                if (
                    skill
                    and len(skill) > 2            # remove very short words
                    and not skill.isdigit()      # remove numbers
                ):
                    skills.append(skill)

    return list(set(skills))

skill_keywords = load_skill_keywords()
print("✅ Loaded skills (first 20):", skill_keywords[:20])
print("✅ Total skills loaded:", len(skill_keywords))

# -------- HELPER FUNCTIONS --------
def ask_ollama(prompt):
    try:
        response = requests.post(
            "http://localhost:11434/api/generate",
            json={
                "model": "llama3",
                "prompt": prompt,
                "stream": False
            }
        )

        return response.json()["response"]

    except Exception as e:
        return "Error generating question"

def ask_ollama_eval(prompt):
    try:
        result = subprocess.run(
            ["ollama", "run", "llama3"],
            input=prompt,
            capture_output=True,
            text=True,
            timeout=60
        )

        output = result.stdout.strip()

        if not output:
            return (
                "Overall Score: 6/10\n"
                "Strengths: Answers show basic understanding.\n"
                "Weaknesses: Lacks depth and examples.\n"
                "Recommendation: Practice structured responses."
            )

        return output

    except Exception as e:
        return (
            "Overall Score: 5/10\n"
            "Strengths: Attempted most questions.\n"
            "Weaknesses: Answers were brief.\n"
            "Recommendation: Improve clarity and confidence."
        )

def extract_text_from_pdf(path):
    text = ""
    doc = fitz.open(path)
    for page in doc:
        text += page.get_text()
    text = re.sub(r"\s+", " ", text)
    return text

# def extract_resume_skills(text):
#     doc = nlp(text)
#     return list(set(
#         token.lemma_.lower()
#         for token in doc
#         if token.lemma_.lower() in skill_keywords
#     ))

def extract_resume_skills(text):
    text = text.lower()
    matched_skills = []

    for skill in skill_keywords:

        # ❗ Skip noisy/invalid skills
        if len(skill) <= 2:
            continue

        # ❗ Skip very generic words
        if skill in ["data", "code", "system", "process", "analysis"]:
            continue

        pattern = r"\b" + re.escape(skill) + r"\b"

        if re.search(pattern, text):
            matched_skills.append(skill)

    print("\n🧠 Clean Resume Skills:", matched_skills)

    return list(set(matched_skills))

def get_interviewer_prompt(job_role, mode):
    if mode.lower() == "hr":
       interviewer = "HR Interviewer"
       rules = """
    ONLY behavioral, situational, and soft-skill questions.
    NO coding, NO algorithms, NO ML, NO technical terms.
    """
    else:
        interviewer = "Technical Interviewer"
        rules = f"{mode} level technical questions for {job_role}"

    return f"""
You are a professional {interviewer}.

JOB ROLE: {job_role}
MODE: {mode}
FOCUS: {rules}

STRICT RULES:
- Ask EXACTLY 10 questions
- Ask ONE question at a time
- Wait for user to submit or skip
- Do NOT give feedback now
- Maintain difficulty strictly
- HR mode must avoid technical questions

Start with Question 1.
"""

def reset_interview(job_role, mode):
    session["interview"] = {
        "job_role": job_role,
        "mode": mode,
        "questions": [],
        "answers": [],
        "timestamps": [],
        "skipped": [],
        "current": 0,
        "completed": False 
    }

# -------- PAGE ROUTES --------

@app.route("/")
def home():
    return render_template("index.html")

@app.route("/about")
def about():
    return render_template("about.html")

@app.route("/dashboard")
def dashboard():
    if "user" not in session:
        flash("Please login to access dashboard", "error")
        return redirect(url_for("login"))

    return render_template("dashboard.html")

@app.route("/interview")
def interview():
    return render_template("interview.html")

@app.route("/login")
def login():
    mode = request.args.get("mode", "login")
    return render_template("login.html", mode=mode)

@app.route("/module1")
def module1():
    return render_template("module1.html")

@app.route("/module2")
def module2():
    return render_template("module2.html")

@app.route("/module3")
def module3():
    return render_template("module3.html")


@app.route("/api/job-roles")
def get_job_roles():
    roles = job_role_df['Job Role'].dropna().unique().tolist()
    return jsonify(roles)

# @app.route("/api/interview/start", methods=["POST"])
# def interview_start():

#     data = request.json
#     job_role = data.get("job_role")
#     mode = data.get("mode")

#     if not job_role:
#         return jsonify({"error": "Job role required"}), 400

#     reset_interview(job_role, mode)

#     prompt = f"""
# You are an expert technical interviewer.

# Generate EXACTLY 10 interview questions.

# STRICT RULES:
# - Output ONLY the questions
# - Do NOT write "Here are the questions"
# - Do NOT write explanations
# - Do NOT write headings
# - Each question must be on a new line
# - Start directly with Question 1

# Job Role:
# {job_role}

# Difficulty:
# {mode}
# """

#     questions_text = ask_ollama(prompt)

#     questions = [
#     q.strip()
#     for q in questions_text.split("\n")
#     if q.strip() and not q.lower().startswith("here")
# ]

#     session["interview"]["questions"] = questions
#     session["interview"]["timestamps"].append(time.time())

#     return jsonify({
#         "question_no": 1,
#         "question": questions[0]
#     })

@app.route("/api/interview/start", methods=["POST"])
def interview_start():
    data = request.json
    job_role = data.get("job_role")
    mode = data.get("mode")

    if not job_role:
        return jsonify({"error": "Job role required"}), 400

    reset_interview(job_role, mode)

    prompt = f"""
You are an expert technical interviewer.

Generate EXACTLY 10 interview questions.

STRICT RULES:
- Output ONLY the questions
- Do NOT write "Here are the questions"
- Do NOT write explanations
- Do NOT write headings
- Each question must be on a new line
- Start directly with Question 1

Job Role:
{job_role}

Difficulty:
{mode}
"""

    questions_text = ask_ollama(prompt)

    questions = [
        q.strip("0123456789.:-) ").strip()
        for q in questions_text.split("\n")
        if q.strip() and not q.lower().startswith("here")
    ]

    questions = [q for q in questions if len(q) > 8]
    questions = questions[:10]

    # ✅ Force exactly 10 questions
    if len(questions) < 10:
        return jsonify({
            "error": f"Ollama generated only {len(questions)} valid questions. Please try again."
        }), 500

    session["interview"]["questions"] = questions
    session["interview"]["timestamps"].append(time.time())
    session.modified = True

    return jsonify({
        "question_no": 1,
        "question": questions[0]
    })

# @app.route("/api/interview/next", methods=["POST"])
# def interview_next():

#     data = request.json
#     action = data.get("action")
#     answer = data.get("answer", "")

#     interview = session.get("interview")

#     if interview.get("completed"):
#         return jsonify({"done": True})

#     now = time.time()
#     time_taken = now - interview["timestamps"][-1]

#     if action == "submit":

#         interview["answers"].append({
#             "text": answer,
#             "time": time_taken
#         })

#     else:
#         interview["skipped"].append(len(interview["answers"]))

#     current_index = len(interview["answers"]) + len(interview["skipped"])

#     if current_index >= 10:

#         interview["completed"] = True
#         session.modified = True

#         return jsonify({"done": True})

#     next_question = interview["questions"][current_index]

#     interview["timestamps"].append(time.time())

#     session.modified = True

#     return jsonify({
#         "question_no": current_index + 1,
#         "question": next_question
#     })

@app.route("/api/interview/next", methods=["POST"])
def interview_next():
    data = request.json
    action = data.get("action")
    answer = data.get("answer", "")

    interview = session.get("interview")

    if not interview:
        return jsonify({"error": "Interview session not found"}), 400

    if interview.get("completed"):
        return jsonify({"done": True})

    if not interview.get("questions"):
        return jsonify({"error": "No interview questions found"}), 400

    now = time.time()
    time_taken = now - interview["timestamps"][-1]

    if action == "submit":
        interview["answers"].append({
            "text": answer,
            "time": time_taken
        })
    else:
        interview["skipped"].append(len(interview["answers"]))

    current_index = len(interview["answers"]) + len(interview["skipped"])
    total_questions = len(interview["questions"])

    if current_index >= total_questions:
        interview["completed"] = True
        session.modified = True
        return jsonify({"done": True})

    next_question = interview["questions"][current_index]

    interview["timestamps"].append(time.time())
    session.modified = True

    return jsonify({
        "question_no": current_index + 1,
        "question": next_question
    })

# def detect_cheating(answers):
#     flags = []

#     # Very fast answers
#     fast = sum(1 for a in answers if a["time"] < 3)
#     if fast >= 3:
#         flags.append("Answered multiple questions too quickly")

#     # Copy-paste suspicion (long text, fast)
#     for a in answers:
#         if len(a["text"].split()) > 80 and a["time"] < 5:
#             flags.append("Possible copy-paste detected")

#     # Repetitive answers
#     for i in range(len(answers)-1):
#         sim = SequenceMatcher(None, answers[i]["text"], answers[i+1]["text"]).ratio()
#         if sim > 0.85:
#             flags.append("Highly repetitive answers detected")

#     return list(set(flags))

def detect_cheating(answers):
    flags = []

    # Very fast answers only if many are too short and too fast
    fast_short = sum(
        1 for a in answers
        if a["time"] < 5 and len(a["text"].split()) < 12
    )
    if fast_short >= 4:
        flags.append("Several answers were submitted unusually quickly.")

    # Copy-paste suspicion only for very long and unrealistically fast responses
    for a in answers:
        if len(a["text"].split()) > 100 and a["time"] < 6:
            flags.append("Possible copy-paste detected in one or more answers.")
            break

    # Repetitive answers only if text is long enough and similarity is very high
    repetitive_count = 0
    for i in range(len(answers) - 1):
        a1 = answers[i]["text"].strip()
        a2 = answers[i + 1]["text"].strip()

        if len(a1.split()) < 8 or len(a2.split()) < 8:
            continue

        sim = SequenceMatcher(None, a1, a2).ratio()
        if sim > 0.93:
            repetitive_count += 1

    if repetitive_count >= 2:
        flags.append("Some answers were highly repetitive.")

    return list(set(flags))

# def calculate_score(answers, skipped):
#     score = 0
#     for a in answers:
#         length = len(a["text"].split())
#         length_score = 10 if 40 <= length <= 120 else max(2, length // 15)
#         time_penalty = 2 if a["time"] < 3 else 0
#         score += max(0, length_score - time_penalty)

#     score -= len(skipped) * 2
#     return max(0, min(100, score))

def calculate_score(answers, skipped):
    total = 0

    for a in answers:
        text = a["text"].strip()
        length = len(text.split())
        time_taken = a["time"]

        score = 0

        # Base score for attempting properly
        if length >= 8:
            score += 4
        elif length >= 4:
            score += 2
        else:
            score += 1

        # Better depth
        if 20 <= length <= 120:
            score += 3
        elif length > 120:
            score += 2

        # Reward structured / thoughtful answers
        if any(word in text.lower() for word in ["because", "for example", "for instance", "using", "approach", "method"]):
            score += 2

        # Penalize only unrealistically fast tiny answers
        if time_taken < 4 and length < 10:
            score -= 2

        # Keep per-answer score within 1–10
        score = max(1, min(10, score))
        total += score

    # Scale to 100
    max_possible = 10 * max(len(answers) + len(skipped), 1)
    score_out_of_100 = (total / max_possible) * 100

    # Small penalty for skipped questions
    score_out_of_100 -= len(skipped) * 3

    return max(0, min(100, round(score_out_of_100)))

@app.route("/api/interview/feedback")
def interview_feedback():

    interview = session["interview"]

    questions = interview["questions"]
    answers = [a["text"] for a in interview["answers"]]

#     prompt = f"""
# You are a strict technical interview evaluator.

# Evaluate each answer based on the question.

# Provide:
# - Score out of 10 for each answer
# - Strengths
# - Weaknesses
# - Suggestions

# Finally provide:
# Overall Score out of 100
# Final Verdict

# Questions:
# {questions}

# Answers:
# {answers}
# """

    prompt = f"""
You are an expert interview evaluator.

Evaluate the candidate's interview performance across ALL answers, not just one answer.

IMPORTANT INSTRUCTIONS:
- Consider the full interview overall
- Judge correctness, clarity, relevance, and technical depth
- Do not focus only on the first answer
- Give realistic scoring, not overly harsh scoring
- If some answers are strong and some are weak, reflect that balance fairly

Return the response EXACTLY in this format:

OVERALL_SCORE: <number out of 100>

STRENGTHS:
- <overall strength 1>
- <overall strength 2>
- <overall strength 3>

WEAKNESSES:
- <overall weakness 1>
- <overall weakness 2>
- <overall weakness 3>

SUGGESTIONS:
- <overall suggestion 1>
- <overall suggestion 2>
- <overall suggestion 3>

FINAL_VERDICT:
- <short overall verdict>

Questions:
{questions}

Answers:
{answers}
"""

#     prompt = f"""
# You are a strict technical interviewer.

# Evaluate the candidate answers.

# IMPORTANT FORMAT:

# Return the response EXACTLY in this format.

# OVERALL_SCORE: <number out of 100>

# STRENGTHS:
# - point
# - point
# - point

# WEAKNESSES:
# - point
# - point

# IMPROVEMENTS:
# - point
# - point

# Questions:
# {questions}

# Answers:
# {answers}
# """

    final_report = ask_ollama(prompt)

    system_score = calculate_score(
    interview["answers"],
    interview["skipped"]
)

# Extract AI score
    import re
    match = re.search(r"OVERALL_SCORE:\s*(\d+)", final_report)
    if match:
        ai_score = int(match.group(1))
    else:
        ai_score = system_score

# Combine scores
    final_score = int((ai_score * 0.75) + (system_score * 0.25))

    cheating_flags = detect_cheating(interview["answers"])

    return jsonify({
        "final_report": final_report,
        "final_score": final_score,
        "cheating_flags": cheating_flags
    })

@app.route("/api/interview/context", methods=["POST"])
def set_interview_context():
    data = request.json

    session["preloaded_interview"] = {
        "job_role": data.get("job_role"),
        "roles": data.get("roles")  # for module 2
    }

    return jsonify({"status": "context set"})

@app.route("/signup", methods=["POST"])
def signup():
    full_name = request.form.get("full_name", "").strip()
    email = request.form.get("email", "").strip().lower()
    password = request.form.get("password", "")
    confirm_password = request.form.get("confirm_password", "")

    if not full_name or not email or not password or not confirm_password:
        flash("Please fill all signup fields.", "error")
        return redirect(url_for("login"))

    if password != confirm_password:
        flash("Passwords do not match.", "error")
        return redirect(url_for("login"))

    existing_user = get_user_by_email(email)
    if existing_user:
        flash("Account already exists. Please log in.", "error")
        return redirect(url_for("login"))

    hashed_password = generate_password_hash(password)
    create_user(full_name, email, hashed_password)

    session["user"] = email
    session["first_name"] = full_name.split()[0]
    session["user_name"] = full_name

    
    return redirect(url_for("home"))

@app.route("/login-user", methods=["POST"])
def login_user():
    email = request.form.get("email", "").strip().lower()
    password = request.form.get("password", "")

    if not email or not password:
        flash("Please enter email and password.", "error")
        return redirect(url_for("login"))

    user = get_user_by_email(email)

    if not user:
        flash("No account found. Please sign up first.", "error")
        return redirect(url_for("login", mode="signup"))

    if not check_password_hash(user["password"], password):
       flash("Incorrect password.", "error")
       return redirect(url_for("login"))

    session["user"] = user["email"]
    session["first_name"] = user["full_name"].split()[0]
    session["user_name"] = user["full_name"]

    return redirect(url_for("home"))

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("home"))

@app.context_processor
def inject_user():
    full_name = session.get("user_name")
    first_name = full_name.split()[0] if full_name else None
    return dict(first_name=first_name)

def normalize_skill(skill):
    skill = skill.lower().strip()

    # remove brackets noise
    skill = re.sub(r"[()]", "", skill)

    # standardize variants
    replacements = {
        "javascript": "js",
        "react.js": "react",
        "node.js": "node",
        "vue.js": "vue",
        "angular.js": "angular"
    }

    return replacements.get(skill, skill.replace(".js", ""))

# -------- API: MODULE 1 --------
# @app.route("/api/skill-gap", methods=["POST"])

# def skill_gap():
#     resume = request.files['resume']
#     job_role = request.form['job_role']
#     threading.Thread(target=run_job_fetcher, args=(job_role,)).start()

#     path = os.path.join(app.config['UPLOAD_FOLDER'], resume.filename)
#     resume.save(path)

#     resume_text = extract_text_from_pdf(path)
#     print("\n📄 Resume Text Preview:\n", resume_text[:500])
#     resume_skills = extract_resume_skills(resume_text)

#     resume_skills = [normalize_skill(s) for s in resume_skills]
#     required_skills = [normalize_skill(s) for s in required_skills]

#     row = job_role_df[job_role_df['Job Role'].str.lower() == job_role.lower()]

#     if row.empty:
#        return jsonify({"error": "Invalid job role"}), 400

#     required_skills = [
#     re.sub(r"[()]", "", s).strip().lower()
#     for s in row.iloc[0]['Technical Skills Needed'].split(',')
#     if s.strip()
# ]
#     print("\n📊 Required Skills:", required_skills)


#     matched = [s for s in required_skills if s in resume_skills]
#     missing = [s for s in required_skills if s not in resume_skills]

#     return jsonify({
#         "job_role": job_role,
#         "matched_skills": matched,
#         "missing_skills": missing
#     })

@app.route("/api/skill-gap", methods=["POST"])
def skill_gap():
    resume = request.files['resume']
    job_role = request.form['job_role']

    threading.Thread(target=run_job_fetcher, args=(job_role,)).start()

    path = os.path.join(app.config['UPLOAD_FOLDER'], resume.filename)
    resume.save(path)

    # 1️⃣ Extract text
    resume_text = extract_text_from_pdf(path)
    print("\n📄 Resume Text Preview:\n", resume_text[:500])

    # 2️⃣ Extract resume skills
    resume_skills = extract_resume_skills(resume_text)

    # 3️⃣ Get job role row
    row = job_role_df[job_role_df['Job Role'].str.lower() == job_role.lower()]

    if row.empty:
        return jsonify({"error": "Invalid job role"}), 400

    # 4️⃣ Get required skills
    required_skills = [
        re.sub(r"[()]", "", s).strip().lower()
        for s in row.iloc[0]['Technical Skills Needed'].split(',')
        if s.strip()
    ]

    print("\n📊 Required Skills:", required_skills)

    # 5️⃣ Normalize BOTH
    resume_skills = [normalize_skill(s) for s in resume_skills]
    required_skills = [normalize_skill(s) for s in required_skills]

    # 6️⃣ Match correctly
    matched = [s for s in required_skills if s in resume_skills]
    missing = [s for s in required_skills if s not in resume_skills]

    return jsonify({
        "job_role": job_role,
        "matched_skills": matched,
        "missing_skills": missing
    })

# -------- API: MODULE 2 --------
@app.route("/api/role-match", methods=["POST"])
def role_match():
    resume = request.files['resume']

    path = os.path.join(app.config['UPLOAD_FOLDER'], resume.filename)
    resume.save(path)

    resume_text = extract_text_from_pdf(path)
    resume_skill_list = extract_resume_skills(resume_text)
    resume_skills = " ".join(resume_skill_list)

    tfidf = TfidfVectorizer()
    X = tfidf.fit_transform(job_role_df['Technical Skills Needed'])
    resume_vec = tfidf.transform([resume_skills])

    scores = cosine_similarity(resume_vec, X).flatten()

    temp_df = job_role_df.copy()
    temp_df["score"] = scores

    top_roles = (
        temp_df.sort_values("score", ascending=False)
        .groupby("Job Role", as_index=False)
        .first()
        .sort_values("score", ascending=False)
        .head(3)
        .copy()
    )

    results = []

    resume_skill_set = set(normalize_skill(s) for s in resume_skill_list)

    for _, row in top_roles.iterrows():
        # required_skills = [
        #     s.strip().lower()
        #     for s in str(row["Technical Skills Needed"]).split(",")
        #     if s.strip()
        # ]
        required_skills = [
    normalize_skill(re.sub(r"[()]", "", s).strip())
    for s in row["Technical Skills Needed"].split(',')
    if s.strip()
]

        matched_skills = [s for s in required_skills if s in resume_skill_set][:3]

        results.append({
            "Job Role": row["Job Role"],
            "score": float(row["score"]),
            "matched_skills": matched_skills
        })

    return jsonify(results)

def extract_job_role_from_jd(jd_text):
    match = re.search(r"role\s*:\s*(.+)", jd_text, re.IGNORECASE)
    if match:
        return match.group(1).strip()
    return None

@app.route("/api/rejection-analysis", methods=["POST"])

def rejection_analysis():
    try:
        rejected_file = request.files.get("rejected_resume")
        selected_file = request.files.get("selected_resume")
        jd_text = request.form.get("job_description", "").strip()

        if not rejected_file or not selected_file or not jd_text:
            return jsonify({"error": "Missing inputs"}), 400

        # Save files
        rej_filename = secure_filename(rejected_file.filename)
        sel_filename = secure_filename(selected_file.filename)

        rej_path = os.path.join(app.config['UPLOAD_FOLDER'], rej_filename)
        sel_path = os.path.join(app.config['UPLOAD_FOLDER'], sel_filename)

        rejected_file.save(rej_path)
        selected_file.save(sel_path)

        # Run ML Analysis
        result = run_analysis(
            rejected_path=rej_path,
            selected_paths=[sel_path],
            jd_text=jd_text
        )

        # Convert Plotly figures to JSON
        radar_json = plotly.io.to_json(
            plot_radar(result["rejected_features"], result["selected_avg"])
        )

        importance_json = plotly.io.to_json(
            plot_feature_importance(result["model"], result["feature_names"])
        )

        shap_json = plotly.io.to_json(
            plot_shap_bar(result["rejected_shap"], result["feature_names"])
        )

        # Remove files
        os.remove(rej_path)
        os.remove(sel_path)
       
        job_role = extract_job_role_from_jd(jd_text)

        return jsonify({
            "job_role": job_role,
            "probability": float(result["prob_selected"] * 100),
            "jd_similarity": float(result["jd_similarity"] * 100),
            "keyword_coverage": float(result["kw_overlap"] * 100),
            "matched_keywords": result["matched_kw"][:20],
            "missing_keywords": result["missing_kw"][:20],
            "reasons": result["reasons"],
            "roadmap": result["roadmap_df"].to_dict(orient="records"),
            "radar_chart": radar_json,
            "importance_chart": importance_json,
            "shap_chart": shap_json,
            "intelligent_insights": result["intelligent_insights"],
            "has_positive_shap": result["has_positive_shap"]
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500

from module3_backend import (
    run_analysis,
    plot_radar,
    plot_feature_importance,
    plot_shap_bar
)

@app.route("/api/test-ollama")
def test_ollama():
    response = ask_ollama(
        "Ask one basic technical interview question for a Data Scientist."
    )
    return jsonify({"ollama_response": response})

if __name__ == "__main__":
    app.run(debug=True)
